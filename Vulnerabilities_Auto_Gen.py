import sys
import docx.shared
import time
import pandas
import numpy as np
from os import path
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from selenium import webdriver
from selenium.webdriver import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from tkinter import filedialog


def get_empty_CVE(records_CVE):
    records_CVE.append([])
    records_CVE[0].append("Отсутствует")
    records_CVE[0].append("Отсутствует")
    return records_CVE

def get_found_CVE(rows, cols, records, driver):
    for r in range(1, int(rows) + 1):
        records.append([])
        for c in range(1, cols + 1):
            value = driver.find_element(By.XPATH, value='//*[@id="TableWithRules"]/table/tbody/tr[' + str(
                r) + ']/td[' + str(c) + ']').text
            records[r - 1].append(value)
    driver.close()
    return records
def vulnerabilities_search_CVE(soft_name):
    try:
        driver = webdriver.Chrome()
    except:
        sys.exit("Возникла ошибка с веб драйвером Chrome")
    driver.get("https://cve.mitre.org/cve/search_cve_list.html")
    print(driver.title)
    elem = driver.find_element(By.NAME, "keyword")
    elem.send_keys(soft_name)
    elem.submit()
    #rows = len(driver.find_elements(By.XPATH, value='//*[@id="TableWithRules"]/table/tbody/tr'))
    rows = driver.find_element(By.XPATH, '// *[ @ id = "CenterPane"] / div[1] / b').text
    print(soft_name)
    print("Найдено уязвимостей: ", rows)
    cols = 2
    records_CVE = []
    if (int(rows) == 0):
        records_CVE = get_empty_CVE(records_CVE)
        driver.close()
        return records_CVE
    else:
        if (int(rows) > 300):
            print ("Количество найденных уязвимостей для " + soft_name + " слишком велико")
            print ("Возможно, не все они в действительности являются уязвимостями данного ПО")
            print ("Полное выполнение программы займёт больше 1 часа")
            flag = input ("Продолжить генерацию перечня уязвимостей для " + soft_name + "? [y/any other letter]: ")
            if (flag == "y"):
                records_CVE = get_found_CVE(rows, cols, records_CVE, driver)
                return records_CVE
            else:
                records_CVE = get_empty_CVE(records_CVE)
                return records_CVE
        else:
            records_CVE = get_found_CVE(rows, cols, records_CVE, driver)
            return records_CVE

def vulnerabilities_search_BDU(records_CVE):
    records_BDU = []
    records_CVE = np.array(records_CVE)
    CVE_identifiers = records_CVE[:, 0]
    if (CVE_identifiers[0] == "Отсутствует"):
        value = "Отсутствует"
        records_BDU.append(value)
        return records_BDU
    try :
        driver = webdriver.Firefox()
    except:
        sys.exit("Возникла ошибка с веб драйвером FireFox")
    print("Сопоставление идентификаторов CVE и ФСТЭК...")
    try:
        driver.get("https://bdu.fstec.ru/vul/")
        time.sleep(5)
        print (driver.title)
        #button_search = driver.find_element(By.XPATH, value='//*[@id="s2id_VulFilterForm_idval"]/a/span[2]')
        #button_submit = driver.find_element(By.XPATH, value='//*[@id="vul-filter-form"]/div[11]/input[2]')
    except BaseException as exc:
        print (exc.__traceback__)
        sys.exit("Возникла ошибка с доступом к сайту ФСТЭК")
    for i in range(len(CVE_identifiers)):
        try:
            WebDriverWait(driver, 40).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="s2id_VulFilterForm_idval"]/a/span[2]'))).click()
            #button_search = driver.find_element(By.XPATH, '//*[@id="s2id_VulFilterForm_idval"]/a/span[2]')
            #driver.execute_script("arguments[0].click();", button_search)
            WebDriverWait(driver, 20).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="s2id_autogen17_search"]'))).send_keys(CVE_identifiers[i])
            time.sleep(3)
            search_id = WebDriverWait(driver, 40).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="select2-results-17"]/li'))).text
            if (search_id == "Совпадений не найдено"):
                value = "Отсутствует"
                time.sleep(2)
                driver.find_element(By.ID, value='select2-drop-mask').click()
            else:
                actions = ActionChains(driver)
                actions.send_keys(Keys.ENTER)
                actions.perform()
                #WebDriverWait(driver, 20).until(
                 #   EC.element_to_be_clickable((By.XPATH, '//*[@id="vul-filter-form"]/div[11]/input[2]'))).click()
                button_submit = driver.find_element(By.XPATH, value='//*[@id="vul-filter-form"]/div[11]/input[2]')
                driver.execute_script("arguments[0].click();", button_submit)
                time.sleep(4)
                #value = driver.find_element(By.XPATH, value='//*[@id="vuls"]/table/tbody/tr/td[1]/h4/a').text
                value = WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="vuls"]/table/tbody/tr/td[1]/h4/a'))).text
                time.sleep(2)
        except BaseException as exc:
            print("Не удалось произвести соответствие уязвимости " + CVE_identifiers[i])
            print (exc)
            value = "ОШИБКА!"
        print(CVE_identifiers[i], " --- ", value)
        print()
        records_BDU.append(value)
    driver.close()
    return records_BDU

def danger_lvl_form(records_CVE):
    print("Формирование уровней опасности уязвимостей...")
    records_CVE = np.array(records_CVE)
    result = []
    CVE_identifiers = records_CVE[:, 0]
    if (CVE_identifiers[0] == "Отсутствует"):
        danger_lvl_text = "Отсутствует"
        result.append(danger_lvl_text)
        return result
    low = "Низкий "
    medium = "Средний "
    high = "Высокий "
    critical = "Критический "
    try:
        driver = webdriver.Chrome()
    except:
        sys.exit("Возникла ошибка с веб драйвером Chrome")
    time.sleep(1)
    for i in range(len(CVE_identifiers)):
        try:
            driver.get("https://nvd.nist.gov/vuln/detail/" + CVE_identifiers[i])
            cvss_vers_btn = driver.find_element(By.ID, 'btn-cvss2')
            cvss_vers_btn.click()
            base_score = driver.find_element(By.XPATH, value = '// *[ @ id = "Vuln2CvssPanel"] / div[1] / div[2] / span / span').text
            cvss_vers = "2.0"
            if (base_score == "N/A"):
                cvss_vers_btn = driver.find_element(By.ID, 'btn-cvss3')
                cvss_vers_btn.click()
                base_score = driver.find_element(By.ID, 'Cvss3CnaCalculatorAnchor').text
                cvss_vers = "3.0"
            base_score = float(base_score[:3])
            danger_lvl_text = "уровень опасности (базовая оценка CVSS " + cvss_vers + " составляет " + str(base_score) + ")"
            if (base_score < 4.0):
                danger_lvl_text = low + danger_lvl_text
            if (base_score >= 4.0 and base_score < 7.0):
                danger_lvl_text = medium + danger_lvl_text
            if (base_score >= 7.0 and base_score < 10.0):
                danger_lvl_text = high + danger_lvl_text
            if (base_score == 10.0):
                danger_lvl_text = critical + danger_lvl_text
            result.append(danger_lvl_text)
        except BaseException as exc:
            driver.close()
            print("Произошла непредвиденная ошибка при формировании уровня опасности уязвимости")
            input_letter = input("Повторить выполнение ещё раз? [y/any other letter]: ")
            if (input_letter == "y"):
                danger_lvl_form(records_CVE)
            else:
                sys.exit(exc)
    driver.close()
    return result

def table_view(table):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(8)
            font.name = 'Arial'

    table.allow_autofit = False
    if (len(table.columns) == 3):
        for cell in table.columns[0].cells:
            cell.width = Inches(0.8976378)

        for cell in table.columns[1].cells:
            cell.width = Inches(3.251969)

        for cell in table.columns[2].cells:
            cell.width = Inches(1.641732)
    else:
        for cell in table.columns[0].cells:
            cell.width = Inches(0.8976378)

        for cell in table.columns[1].cells:
            cell.width = Inches(0.8976378)

        for cell in table.columns[2].cells:
            cell.width = Inches(3.251969)

        for cell in table.columns[3].cells:
            cell.width = Inches(1.641732)

    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#FFC000"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#FFC000"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

    shading_elm_3 = parse_xml(r'<w:shd {} w:fill="#FFC000"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_3)

    if (len(table.columns) > 3):
        shading_elm_4 = parse_xml(r'<w:shd {} w:fill="#FFC000"/>'.format(nsdecls('w')))
        table.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_4)

def save_doc(doc_path, document):
    document.add_page_break()
    document.save(doc_path)
    print("Документ успешно сохранён")
    print()

def create_table_CVE(doc_path, document, records_CVE):
    danger_lvl_text = np.array(danger_lvl_form(records_CVE))
    records_CVE = np.column_stack((records_CVE, danger_lvl_text))
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Идентификатор CVE'
    hdr_cells[1].text = 'Описание уязвимостей'
    hdr_cells[2].text = 'Уровень опасности уязвимости'
    for id_CVE, desc, lvl in records_CVE:
        row_cells = table.add_row().cells
        row_cells[0].text = id_CVE
        row_cells[1].text = desc
        row_cells[2].text = lvl

    table_view(table)
    save_doc(doc_path, document)
    return (records_CVE)

def create_table_with_BDU(doc_path, document, records_CVE):
    try:
        records_BDU = np.array(vulnerabilities_search_BDU(records_CVE))
    except BaseException as exc:
        print("Произошла непредвиденная ошибка при формировании идентификаторов ФСТЭК")
        input_letter = input(
            "Повторить попытку ещё раз? (При отказе сформируется таблица без идентификаторов ФСТЭК) [y/any other letter]: ")
        if (input_letter == "y"):
            records_BDU = np.array(vulnerabilities_search_BDU(records_CVE))
        else:
            create_table_CVE(doc_path, document, records_CVE)
    records_CVE = np.insert(records_CVE, 1, records_BDU, axis=1)
    danger_lvl_text = np.array(danger_lvl_form(records_CVE))
    records_CVE = np.column_stack((records_CVE, danger_lvl_text))
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Идентификатор CVE'
    hdr_cells[1].text = 'Идентификатор ФСТЭК'
    hdr_cells[2].text = 'Описание уязвимостей'
    hdr_cells[3].text = 'Уровень опасности уязвимости'
    for id_CVE, id_BDU, desc, lvl in records_CVE:
        row_cells = table.add_row().cells
        row_cells[0].text = id_CVE
        row_cells[1].text = id_BDU
        row_cells[2].text = desc
        row_cells[3].text = lvl

    table_view(table)
    save_doc(doc_path, document)
    return (records_CVE)

def init_doc(doc_path, soft_name):
    document = Document(doc_path)
    head_two = document.add_heading('', 2)
    text_head_two = head_two.add_run(soft_name)
    text_head_two.font.name = 'Arial'
    text_head_two.font.size = Pt(12)
    text_head_two.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    text_head_two.font.all_caps = True

    records_CVE = np.array(vulnerabilities_search_CVE(soft_name))
    if (records_CVE[0,0] == "Отсутствует"):
        records_CVE = create_table_with_BDU(doc_path, document, records_CVE)
        return records_CVE
    vul_count = len(records_CVE)
    waiting_time = int((vul_count * 11) / 60)
    if (waiting_time < 1):
        waiting_time = 1
    print("Примерное время ожидания для формирования идентификаторов ФСТЭК: ", waiting_time, " минут(а/ы)")
    records_CVE = create_table_with_BDU(doc_path, document, records_CVE)
    #records_CVE = create_table_CVE(doc_path, document, records_CVE)
    return (records_CVE)

def set_doc_path():
    doc_path = filedialog.askopenfilename(title="Выбор документа (docx)", defaultextension="docx")
    full_name = path.basename(doc_path)
    doc_ext = path.splitext(full_name)[1]
    if (doc_ext != ".docx"):
        sys.exit("Только файл с расширением .docx")
    else:
        return doc_path

def read_txt_file():
    txt_path = filedialog.askopenfilename(title="Выбор текстового файла (txt)", defaultextension="txt")
    base_name = path.basename(txt_path)
    txt_ext = path.splitext(base_name)[1]
    if (txt_ext != ".txt"):
        sys.exit("Только файл с расширением .txt")
    else:
        file_txt = open(txt_path, "r")
        lines = file_txt.readlines()
        return lines

def copy_doc(doc_path):
    document = Document()
    dir_name = path.dirname(doc_path)
    base_name = path.basename(doc_path)
    doc_name = path.splitext(base_name)[0]
    doc_name = doc_name + '_AUTO_GEN.docx'
    new_doc_path = dir_name + "/" + doc_name
    document.save(new_doc_path)
    return new_doc_path

def read_xlsx():
    xl_path = filedialog.askopenfilename(title="Выбор файла Excel (xlsx)", defaultextension="xlsx")
    base_name = path.basename(xl_path)
    xl_ext = path.splitext(base_name)[1]
    if (xl_ext != ".xlsx"):
        sys.exit("Только файл с расширением .xlsx")
    else:
        xl_file = pandas.read_excel(xl_path)
        xl_file = xl_file.dropna()
        soft_names = xl_file["Установленное ПО"].tolist()
        return soft_names


def old_txt_main():
    soft_names = read_txt_file()
    doc_path = set_doc_path()
    doc_path = copy_doc(doc_path)
    for soft_name in soft_names:
        soft_name = soft_name.strip()
        init_doc(doc_path, soft_name)

def main():
    soft_names = read_xlsx()
    doc_path = set_doc_path()
    doc_path = copy_doc(doc_path)
    for soft_name in soft_names:
        try:
            init_doc(doc_path, str(soft_name))
        except:
            print ("Произошла непредвиденная непрошибка при работе с " + soft_name + ". Таблица с " + soft_name + " будет пропущена.")
            continue

if __name__ == '__main__':
    main()
    input("Press Enter to exit")
